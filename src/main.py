import datetime
import base64
import io
import os
import shutil
from dataclasses import asdict
from datetime import timedelta
from hashlib import pbkdf2_hmac
from secrets import token_hex
from pathlib import Path
from fastapi import BackgroundTasks, Body, Depends
from fastapi import FastAPI, HTTPException, Query, Depends, UploadFile, File, Form
from fastapi.responses import FileResponse, JSONResponse, Response
from secrets import token_hex
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.exc import IntegrityError

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, PageBreak, KeepTogether, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document as DocxDocument
from docx.shared import Inches as DocxInches
from src.infrastructure.repositories import DocumentoRepository



#IMPORTACIONES DE SEGURIDAD
from src.infrastructure.security import verificar_password, crear_access_token, obtener_usuario_actual, verificar_rol, generar_hash_password
from src.interfaces.schemas import LoginRequest, TokenResponse, DocumentoRead
#############################


from src.application.use_cases import OtService, RequirementService, TraceabilityService, UseCaseService
from src.infrastructure.db import init_admin_db, init_db
from src.infrastructure.repositories import (
    AuditLogRepository,
    OtAtmRepository,
    OtRepository,
    RequirementRepository,
    TraceabilityRepository,
    UseCaseRepository,
    UserRepository,
    RegionRepository,
    DocumentoRepository,
)
from src.interfaces.schemas import (
    AuditLogCreate,
    AuditLogRead,
    OtCreate,
    OtEstadoUpdate,
    OtRead,
    OtUpdate,
    RequirementCreate,
    RequirementRead,
    RequirementStatusUpdate,
    TraceabilityRead,
    UseCaseCreate,
    UseCaseRead,
    UserCreate,
    UserRead,
    RegionRead,
)


app = FastAPI(title="SSEI API", version="0.3.0")

# Instanciamos el repositorio al inicio de src/main.py
documento_repository = DocumentoRepository()

# Directorio raíz del almacenamiento de archivos de la biblioteca
BIBLIOTECA_UPLOAD_DIR = Path("data") / "biblioteca"

def ejecutar_limpieza_papelera_30_dias() -> None:
    """
    Tarea en segundo plano que busca documentos que lleven más de 30 días en la papelera,
    los elimina físicamente del disco y posteriormente borra el registro de la Base de Datos.
    """
    try:
        # Obtenemos todos los registros expirados (más de 30 días)
        documentos_expirados = documento_repository.get_expired_documents(age_days=30)
        
        for doc in documentos_expirados:
            # Construimos la ruta del archivo físico en el servidor
            ruta_archivo = BIBLIOTECA_UPLOAD_DIR / doc.nombre_sistema
            
            # 1. Borramos el archivo del almacenamiento físico si existe
            if ruta_archivo.exists():
                try:
                    os.remove(ruta_archivo)
                    print(f"Limpieza Física: Archivo removido del disco: {ruta_archivo}")
                except Exception as file_err:
                    print(f"Alerta: No se pudo eliminar el archivo físico {ruta_archivo}. Detalle: {file_err}")
            
            # 2. Borramos permanentemente el metadato de SQLite de manera segura
            documento_repository.permanent_delete(doc.id)
            print(f"Limpieza Base de Datos: Registro removido para ID {doc.id}")
            
    except Exception as exc:
        print(f"Error crítico durante la limpieza automática de la biblioteca: {exc}")
# CORS permite conexiones desde app mobile y panel admin durante desarrollo.
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

requirement_repository = RequirementRepository()
use_case_repository = UseCaseRepository()
traceability_repository = TraceabilityRepository(
    requirement_repository=requirement_repository,
    use_case_repository=use_case_repository,
)
ot_repository = OtRepository()
ot_atm_repository = OtAtmRepository()
user_repository = UserRepository()
audit_log_repository = AuditLogRepository()
region_repository = RegionRepository()
requirement_service = RequirementService(repository=requirement_repository)
use_case_service = UseCaseService(repository=use_case_repository)
traceability_service = TraceabilityService(repository=traceability_repository)
ot_service = OtService(ot_repo=ot_repository, atm_repo=ot_atm_repository)


    
@app.on_event("startup")
def startup_event() -> None:
    # 1. Crea la carpeta de almacenamiento de la biblioteca si no existe
    BIBLIOTECA_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    
    # 2. Ejecutar tareas de inicio clásicas
    init_db()
    init_admin_db()
    seed_admin()
    seed_tecnicos()
    seed_regiones()
    
    # 3. Disparar limpieza automática de archivos expirados en la papelera
    print("Iniciando depuración pasiva de la papelera de reciclaje de documentos...")
    ejecutar_limpieza_papelera_30_dias()


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


# ---------------------------------------------------------------------------
# Requirements
# ---------------------------------------------------------------------------

@app.get("/requirements", response_model=list[RequirementRead])
def list_requirements() -> list[RequirementRead]:
    return [RequirementRead(**asdict(item)) for item in requirement_service.list_requirements()]


@app.post("/requirements", response_model=RequirementRead, status_code=201)
def create_requirement(payload: RequirementCreate) -> RequirementRead:
    requirement = requirement_service.create_requirement(
        title=payload.title,
        description=payload.description,
        priority=payload.priority,
    )
    return RequirementRead(**asdict(requirement))


@app.patch("/requirements/{requirement_id}/status", response_model=RequirementRead)
def update_requirement_status(requirement_id: int, payload: RequirementStatusUpdate) -> RequirementRead:
    requirement = requirement_service.update_requirement_status(requirement_id=requirement_id, status=payload.status)
    if requirement is None:
        raise HTTPException(status_code=404, detail="Requirement not found")
    return RequirementRead(**asdict(requirement))


@app.get("/use-cases", response_model=list[UseCaseRead])
def list_use_cases() -> list[UseCaseRead]:
    return [UseCaseRead(**asdict(item)) for item in use_case_service.list_use_cases()]


@app.post("/use-cases", response_model=UseCaseRead, status_code=201)
def create_use_case(payload: UseCaseCreate) -> UseCaseRead:
    try:
        use_case = use_case_service.create_use_case(
            code=payload.code,
            name=payload.name,
            description=payload.description,
        )
    except IntegrityError as exc:
        raise HTTPException(status_code=409, detail="Use case code already exists") from exc

    return UseCaseRead(**asdict(use_case))


@app.get("/traceability", response_model=list[TraceabilityRead])
def list_traceability() -> list[TraceabilityRead]:
    return [TraceabilityRead(**asdict(item)) for item in traceability_service.list_traces()]


@app.get("/requirements/{requirement_id}/use-cases", response_model=list[UseCaseRead])
def list_use_cases_by_requirement(requirement_id: int) -> list[UseCaseRead]:
    use_cases = traceability_service.list_use_cases_for_requirement(requirement_id=requirement_id)
    if use_cases is None:
        raise HTTPException(status_code=404, detail="Requirement not found")

    return [UseCaseRead(**asdict(item)) for item in use_cases]


@app.post(
    "/requirements/{requirement_id}/use-cases/{use_case_id}",
    response_model=TraceabilityRead,
    status_code=201,
)
def link_requirement_use_case(requirement_id: int, use_case_id: int) -> TraceabilityRead:
    try:
        trace = traceability_service.link_requirement_to_use_case(
            requirement_id=requirement_id,
            use_case_id=use_case_id,
        )
    except ValueError as exc:
        raise HTTPException(status_code=409, detail="Traceability link already exists") from exc

    if trace is None:
        raise HTTPException(status_code=404, detail="Requirement or use case not found")

    return TraceabilityRead(**asdict(trace))


# ---------------------------------------------------------------------------
# OTs - Mobile
# ---------------------------------------------------------------------------

def _ot_to_read(ot) -> OtRead:
    return OtRead(
        id=ot.id,
        banco=ot.banco,
        estado=ot.estado,
        fecha_creacion=ot.fecha_creacion,
        hora_programada=ot.hora_programada,
        region=ot.region, 
        comuna=ot.comuna,
        direccion=ot.direccion,
        tecnico_id=ot.tecnico_id,
        nombre_tecnico=ot.nombre_tecnico,
        nombre_etv=ot.nombre_etv,
        nombre_alarma=ot.nombre_alarma,
        origen_servidor=ot.origen_servidor,
        atms=[
            {
                "id": a.id,
                "ot_id": a.ot_id,
                "etiqueta": a.etiqueta,
                "tipo_servicio": a.tipo_servicio,
                "numero_atm": a.numero_atm,
                "serie_cajero": a.serie_cajero,
                "serie_mmbb": a.serie_mmbb,
                "detalles_servicio": a.detalles_servicio,
                "observaciones": a.observaciones,
            }
            for a in ot.atms
        ],
    )

@app.get("/ots", response_model=list[OtRead])
def list_ots(
    estado: str | None = Query(default=None),
    tecnico_id: int | None = Query(default=None), # Captura el parámetro de consulta opcional
    token_data: dict = Depends(obtener_usuario_actual)
) -> list[OtRead]:
    rol = token_data.get("rol")
    user_id = token_data.get("sub")

    # Si es técnico, forzamos que filtre por su id del token, sin importar lo que pida el query param
    if rol == "tecnico" and user_id:
        tecnico_id_filtro = int(user_id)
    else:
        # Si es admin/coordinador, puede filtrar por el query param 'tecnico_id' si se envía
        tecnico_id_filtro = tecnico_id

    all_ots = ot_service.list_ots(estado=estado)

    if tecnico_id_filtro is not None:
        return [_ot_to_read(ot) for ot in all_ots if ot.tecnico_id == tecnico_id_filtro]

    return [_ot_to_read(ot) for ot in all_ots]

@app.post("/ots", response_model=OtRead, status_code=201)
def create_ot(
    payload: OtCreate,
    token_data: dict = Depends(verificar_rol(["admin", "coordinador"]))  # Solo Admin o Coordinador
) -> OtRead:
    ot = ot_service.create_ot(
        banco=payload.banco,
        region=payload.region,
        comuna=payload.comuna,
        direccion=payload.direccion,
        hora_programada=payload.hora_programada,
        tecnico_id=payload.tecnico_id,
        nombre_tecnico=payload.nombre_tecnico,
        nombre_etv=payload.nombre_etv,
        nombre_alarma=payload.nombre_alarma,
        atms=[a.model_dump() for a in payload.atms],
    )
    return _ot_to_read(ot)


@app.get("/ots/{ot_id}", response_model=OtRead)
def get_ot(
    ot_id: int,
    token_data: dict = Depends(obtener_usuario_actual)  # Requiere login válido
) -> OtRead:
    ot = ot_service.get_ot(ot_id)
    if ot is None:
        raise HTTPException(status_code=404, detail="OT no encontrada")
    return _ot_to_read(ot)



@app.put("/ots/{ot_id}", response_model=OtRead)
def update_ot(
    ot_id: int,
    payload: OtUpdate,
    token_data: dict = Depends(verificar_rol(["admin", "coordinador"]))  # Solo Admin o Coordinador
) -> OtRead:
    ot = ot_service.update_ot(
        ot_id=ot_id,
        banco=payload.banco,
        region=payload.region,
        comuna=payload.comuna,
        direccion=payload.direccion,
        hora_programada=payload.hora_programada,
        tecnico_id=payload.tecnico_id,
        nombre_tecnico=payload.nombre_tecnico,
        nombre_etv=payload.nombre_etv,
        nombre_alarma=payload.nombre_alarma,
        atms=[a.model_dump() for a in payload.atms] if payload.atms is not None else None,
    )
    if ot is None:
        raise HTTPException(status_code=404, detail="OT no encontrada")
    return _ot_to_read(ot)


@app.patch("/ots/{ot_id}/estado", response_model=OtRead)
def update_ot_estado(
    ot_id: int,
    payload: OtEstadoUpdate,
    token_data: dict = Depends(verificar_rol(["admin", "coordinador", "tecnico"]))  # Solo Admin o Coordinador
) -> OtRead:
    ot = ot_service.update_estado(ot_id=ot_id, estado=payload.estado)
    if ot is None:
        raise HTTPException(status_code=404, detail="OT no encontrada")
    return _ot_to_read(ot)


@app.delete("/ots/{ot_id}", status_code=204)
def delete_ot(
    ot_id: int,
    token_data: dict = Depends(verificar_rol(["admin", "coordinador"]))  # Solo Admin o Coordinador
) -> None:
    if not ot_service.delete_ot(ot_id):
        raise HTTPException(status_code=404, detail="OT no encontrada")

# ---------------------------------------------------------------------------
# Admin foundations: users and audit logs
# ---------------------------------------------------------------------------

@app.get("/admin/users", response_model=list[UserRead])
def list_users() -> list[UserRead]:
    users = user_repository.list_all()
    return [
        UserRead(
            id=user.id,
            email=user.email,
            full_name=user.full_name,
            role=user.role,
            is_active=user.is_active,
            created_at=user.created_at,
        )
        for user in users
    ]


@app.post("/admin/users", response_model=UserRead, status_code=201)
def create_user(payload: UserCreate) -> UserRead:
    salt = token_hex(16)
    password_hash = pbkdf2_hmac("sha256", payload.password.encode("utf-8"), bytes.fromhex(salt), 390000).hex()
    stored_hash = f"pbkdf2_sha256${salt}${password_hash}"
    try:
        user = user_repository.create(
            email=payload.email,
            hashed_password=stored_hash,
            full_name=payload.full_name,
            role=payload.role,
            is_active=payload.is_active,
        )
    except IntegrityError as exc:
        raise HTTPException(status_code=409, detail="User email already exists") from exc

    return UserRead(
        id=user.id,
        email=user.email,
        full_name=user.full_name,
        role=user.role,
        is_active=user.is_active,
        created_at=user.created_at,
    )
@app.delete("/admin/users/{user_id}", status_code=204)
def delete_user(user_id: int) -> None:
    try:
        if not user_repository.delete(user_id):
            raise HTTPException(status_code=404, detail="Usuario no encontrado")
    except IntegrityError as exc:
        raise HTTPException(
            status_code=409, 
            detail="No se puede borrar el usuario porque tiene registros históricos o auditorías asociadas. Considere desactivarlo."
        ) from exc


@app.patch("/admin/users/{user_id}/status", response_model=UserRead)
def toggle_user_status(user_id: int, is_active: bool = Query(...)) -> UserRead:
    user = user_repository.toggle_active(user_id, is_active)
    if not user:
        raise HTTPException(status_code=404, detail="Usuario no encontrado")
    return UserRead(
        id=user.id,
        email=user.email,
        full_name=user.full_name,
        role=user.role,
        is_active=user.is_active,
        created_at=user.created_at
    )
# ---------------------------------------------------------------------------
# Autenticación Endpoints
# ---------------------------------------------------------------------------

@app.post("/auth/login", response_model=TokenResponse)
def login(payload: LoginRequest) -> TokenResponse:
    # 1. Buscar al usuario en la base de datos por email
    user = user_repository.get_by_email(payload.email)
    if not user:
        raise HTTPException(status_code=401, detail="Credenciales incorrectas")

    # 2. Validar que la cuenta esté activa para iniciar sesión
    if not user.is_active:
        raise HTTPException(status_code=403, detail="Usuario desactivado. Contacte al administrador.")

    # 3. Validar coincidencia de la contraseña plana contra el Hash PBKDF2
    is_valid = verificar_password(payload.password, user.hashed_password)
    if not is_valid:
        raise HTTPException(status_code=401, detail="Credenciales incorrectas")

    # 4. Generar Claims para el Payload JWT
    token_payload = {
        "sub": user.email,
        "id": user.id,
        "name": user.full_name,
        "role": user.role
    }

    # 5. Firmar JWT
    token = crear_access_token(data=token_payload)

    # 6. Preparar lectura de los datos públicos del usuario
    user_read_data = UserRead(
        id=user.id,
        email=user.email,
        full_name=user.full_name,
        role=user.role,
        is_active=user.is_active,
        created_at=user.created_at
    )

    return TokenResponse(
        access_token=token,
        token_type="bearer",
        user=user_read_data
    )

@app.get("/admin/audit-logs", response_model=list[AuditLogRead])
def list_audit_logs() -> list[AuditLogRead]:
    logs = audit_log_repository.list_all()
    return [AuditLogRead(**asdict(log)) for log in logs]


@app.post("/admin/audit-logs", response_model=AuditLogRead, status_code=201)
def create_audit_log(payload: AuditLogCreate) -> AuditLogRead:
    log = audit_log_repository.create(
        user_id=payload.user_id,
        action=payload.action,
        details=payload.details,
    )
    return AuditLogRead(**asdict(log))

def seed_tecnicos() -> None:
    """Registra técnicos por defecto en la base de datos si no existen."""
    tecnicos_por_defecto = [
        {"email": "rodolfo@ssei.cl", "full_name": "Rodolfo Carreño", "role": "tecnico"},
        {"email": "juan@ssei.cl", "full_name": "Juan Albornoz", "role": "tecnico"},
        {"email": "pedro@ssei.cl", "full_name": "Pedro Berrios", "role": "tecnico"},
    ]
    
    for tec in tecnicos_por_defecto:
        try:
            # 🔑 Generación limpia y alineada de la contraseña usando el hash oficial del sistema
            stored_hash = generar_hash_password("password123")
            
            user_repository.create(
                email=tec["email"],
                hashed_password=stored_hash,
                full_name=tec["full_name"],
                role=tec["role"],
                is_active=True
            )
            print(f"Técnico registrado exitosamente: {tec['full_name']}")
        except IntegrityError:
            pass
        
def seed_admin() -> None:
    """Registra el Administrador por defecto si no existe en la base de datos."""
    admin_default = {
        "email": "admin@ssei.cl", 
        "full_name": "Administrador Principal", 
        "role": "admin"
    }
    try:
        stored_hash = generar_hash_password("Asd12345*")
        user_repository.create(
            email=admin_default["email"],
            hashed_password=stored_hash,
            full_name=admin_default["full_name"],
            role=admin_default["role"],
            is_active=True
        )
        print(f"Administrador semilla registrado: {admin_default['email']}")
    except IntegrityError:
        pass

def seed_regiones() -> None:
    if region_repository.count() == 0:
        regiones_chile = [
            'Región de Arica y Parinacota',
            'Región de Tarapacá',
            'Región de Antofagasta',
            'Región de Atacama',
            'Región de Coquimbo',
            'Región de Valparaíso',
            'Región Metropolitana',
            'Región de O’Higgins',
            'Región del Maule',
            'Región de Ñuble',
            'Región del Biobío',
            'Región de La Araucanía',
            'Región de Los Ríos',
            'Región de Los Lagos',
            'Región de Aysén',
            'Región de Magallanes'
        ]
        region_repository.bulk_create(regiones_chile)

@app.on_event("startup")
def startup_event() -> None:
    init_db()
    init_admin_db()
    seed_admin()  # <-- Añadir esta llamada
    seed_tecnicos()
    seed_regiones()  # <-- Añadir esta llamada
    
@app.get("/regiones", response_model=list[RegionRead])
def list_regiones() -> list[RegionRead]:
    return [RegionRead(id=r.id, nombre=r.nombre) for r in region_repository.list_all()]

# ---------------------------------------------------------------------------
# Biblioteca Endpoints
# ---------------------------------------------------------------------------
from src.interfaces.schemas import DocumentoRead

@app.post("/biblioteca/upload", response_model=DocumentoRead, status_code=201)
def upload_documento(
    categoria: str = Form(...),
    banco: str | None = Form(None),
    numero_atm: str | None = Form(None),
    file: UploadFile = File(...),
    token_data: dict = Depends(verificar_rol(["admin", "coordinador", "tecnico"]))
):
    try:
        # 1. Definir rutas relativas seguras basadas en Año/Mes
        ahora = datetime.datetime.now()
        anio = ahora.strftime("%Y")
        mes = ahora.strftime("%m")
        
        # Crear carpeta física de disco correspondiente
        subcarpeta_fecha = Path(anio) / mes
        directorio_destino = BIBLIOTECA_UPLOAD_DIR / subcarpeta_fecha
        directorio_destino.mkdir(parents=True, exist_ok=True)
        
        # 2. Renombrar archivo usando un UUID o token_hex para evitar colisiones
        ext = os.path.splitext(file.filename)[1]
        nombre_renombrado = f"{token_hex(16)}{ext}"
        ruta_archivo_fisico = directorio_destino / nombre_renombrado
        
        # 3. Guardar el archivo físico en el almacenamiento de disco
        with ruta_archivo_fisico.open("wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # 4. Obtener el peso real del archivo escrito en disco en Bytes
        peso_bytes = ruta_archivo_fisico.stat().st_size
        
        # 5. Registrar los metadatos mapeados en la Base de Datos SQLite
        ruta_relativa_sistema = str(subcarpeta_fecha / nombre_renombrado).replace("\\", "/")
        creador_id = token_data.get("id")
        
        doc = documento_repository.create(
            nombre_original=file.filename,
            nombre_sistema=ruta_relativa_sistema,
            peso_bytes=peso_bytes,
            mimetype=file.content_type or "application/octet-stream",
            categoria=categoria,
            banco=banco if banco else None,
            numero_atm=numero_atm if numero_atm else None,
            subido_por_id=creador_id
        )
        
        return DocumentoRead(
            id=doc.id,
            nombre_original=doc.nombre_original,
            nombre_sistema=doc.nombre_sistema,
            peso_bytes=doc.peso_bytes,
            mimetype=doc.mimetype,
            categoria=doc.categoria,
            banco=doc.banco,
            numero_atm=doc.numero_atm,
            subido_por_id=doc.subido_por_id,
            created_at=doc.created_at
        )

    except Exception as exc:
        print(f"Error crítico en subida de biblioteca: {exc}")
        return JSONResponse(
            status_code=500,
            content={"detail": f"Error interno del servidor al procesar el archivo: {str(exc)}"},
            headers={
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "*",
                "Access-Control-Allow-Headers": "*"
            }
        )


@app.get("/biblioteca/documentos", response_model=list[DocumentoRead])
def list_documentos_activos(
    token_data: dict = Depends(obtener_usuario_actual) # Todo usuario autenticado puede visualizar
) -> list[DocumentoRead]:
    docs = documento_repository.list_all_active()
    return [DocumentoRead(
        id=d.id,
        nombre_original=d.nombre_original,
        nombre_sistema=d.nombre_sistema,
        peso_bytes=d.peso_bytes,
        mimetype=d.mimetype,
        categoria=d.categoria,
        banco=d.banco,
        numero_atm=d.numero_atm,
        subido_por_id=d.subido_por_id,
        created_at=d.created_at,
        deleted_at=d.deleted_at,
        deleted_by_id=d.deleted_by_id
    ) for d in docs]


@app.get("/biblioteca/papelera", response_model=list[DocumentoRead])
def list_documentos_papelera(
    token_data: dict = Depends(verificar_rol(["admin"])) # Solo el Administrador ve la Papelera de Reciclaje
) -> list[DocumentoRead]:
    docs = documento_repository.list_trash()
    return [DocumentoRead(
        id=d.id,
        nombre_original=d.nombre_original,
        nombre_sistema=d.nombre_sistema,
        peso_bytes=d.peso_bytes,
        mimetype=d.mimetype,
        categoria=d.categoria,
        banco=d.banco,
        numero_atm=d.numero_atm,
        subido_por_id=d.subido_por_id,
        created_at=d.created_at,
        deleted_at=d.deleted_at,
        deleted_by_id=d.deleted_by_id
    ) for d in docs]


@app.get("/biblioteca/download/{doc_id}")
def download_documento(
    doc_id: int,
    token_data: dict = Depends(obtener_usuario_actual) # Descarga segura para todo usuario autenticado
) -> FileResponse:
    doc = documento_repository.get_by_id(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado o eliminado.")
        
    ruta_archivo = BIBLIOTECA_UPLOAD_DIR / doc.nombre_sistema
    if not ruta_archivo.exists():
        raise HTTPException(status_code=404, detail="El archivo físico ya no se encuentra en el disco del servidor.")
        
    # Retorna la descarga binaria nativa renombrándola al nombre original del cliente
    return FileResponse(
        path=ruta_archivo,
        media_type=doc.mimetype,
        filename=doc.nombre_original
    )


@app.patch("/biblioteca/{doc_id}/trash", response_model=DocumentoRead)
def soft_delete_documento(
    doc_id: int,
    token_data: dict = Depends(verificar_rol(["admin"])) # Solo el Administrador puede enviar a la Papelera
) -> DocumentoRead:
    usuario_id = token_data.get("id")
    doc = documento_repository.soft_delete(doc_id, user_id=usuario_id)
    if not doc:
         raise HTTPException(status_code=404, detail="Documento no encontrado.")
    return DocumentoRead(
        id=doc.id,
        nombre_original=doc.nombre_original,
        nombre_sistema=doc.nombre_sistema,
        peso_bytes=doc.peso_bytes,
        mimetype=doc.mimetype,
        categoria=doc.categoria,
        banco=doc.banco,
        numero_atm=doc.numero_atm,
        subido_por_id=doc.subido_por_id,
        created_at=doc.created_at,
        deleted_at=doc.deleted_at,
        deleted_by_id=doc.deleted_by_id
    )


@app.patch("/biblioteca/{doc_id}/restore", response_model=DocumentoRead)
def restore_documento(
    doc_id: int,
    token_data: dict = Depends(verificar_rol(["admin"])) # Solo el Administrador puede restaurar la papelera
) -> DocumentoRead:
    doc = documento_repository.restore(doc_id)
    if not doc:
         raise HTTPException(status_code=404, detail="Documento no encontrado.")
    return DocumentoRead(
        id=doc.id,
        nombre_original=doc.nombre_original,
        nombre_sistema=doc.nombre_sistema,
        peso_bytes=doc.peso_bytes,
        mimetype=doc.mimetype,
        categoria=doc.categoria,
        banco=doc.banco,
        numero_atm=doc.numero_atm,
        subido_por_id=doc.subido_por_id,
        created_at=doc.created_at,
        deleted_at=doc.deleted_at,
        deleted_by_id=doc.deleted_by_id
    )

@app.post("/biblioteca/generar-informe-completo")
def generar_y_guardar_informe_completo(
    payload: dict = Body(...),
    token_data: dict = Depends(verificar_rol(["admin", "coordinador"]))
):
    """
    Especificación del Informe:
    - Página 1: Imagen del Cliente + Logo Empresa + Título 'Informe Técnico'
    - Página 2: Tipo de Solicitud + Imagen de Servicio (proporcionada por Coordinador)
    - Detalle Técnico: Listado enumerado del informe paso a paso
    - Informe Fotográfico: Título descriptivo + fotos del paquete ZIP
    - Orden de trabajo: Imagen de la OT firmada adjunta
    
    Genera el PDF obligatorio para la biblioteca organizada del ATM y retorna 
    el archivo seleccionado (.pdf o .docx) para su descarga local.
    """
    try:
        dto_atm = payload.get("numeroAtm", "ATM_PROCESADO")
        dto_banco = payload.get("banco", "SSEI_BANCO")
        solicitud = payload.get("tipoSolicitud", "Mantenimiento Preventivo / Correctivo")
        comentario_general = payload.get("comentarioGeneral", "")
        formato = payload.get("formato", "pdf")  # 'pdf' o 'docx'
        
        # Elementos dinámicos en base64
        cliente_logo_b64 = payload.get("clienteLogoBase64")  # Logo Cliente (Varia en Portada)
        servicio_img_b64 = payload.get("servicioImgBase64")    # Imagen del Servicio
        ot_img_b64 = payload.get("otImgBase64")                # Imagen del PDF de la OT
        
        detalle_tecnico_pasos = payload.get("detalleTecnico", []) # List[] de strings
        secciones_fotos = payload.get("fotosAsignadas", [])       # List[] de {titulo, imagen_base64}
        
        ahora = datetime.datetime.now()
        anio = ahora.strftime("%Y")
        mes = ahora.strftime("%m")
        
        # 1. Definir subcarpeta de guardado de biblioteca agrupado por ATM
        subcarpeta_destino = BIBLIOTECA_UPLOAD_DIR / f"{dto_atm}" / anio / mes
        subcarpeta_destino.mkdir(parents=True, exist_ok=True)
        
        token_id = token_hex(8)
        nombre_pdf = f"INFORME_{dto_atm}_{ahora.strftime('%Y%m%d')}_{token_id}.pdf"
        nombre_docx = f"INFORME_{dto_atm}_{ahora.strftime('%Y%m%d')}_{token_id}.docx"
        
        ruta_pdf = subcarpeta_destino / nombre_pdf
        ruta_docx = subcarpeta_destino / nombre_docx

        # Auxiliar para convertir Base64 a Bytes para ReportLab y Docx
        def procesar_b64_img(b64_string):
            if not b64_string or "," not in b64_string:
                return None
            try:
                header, encoded = b64_string.split(",", 1)
                img_data = base64.b64decode(encoded)
                return io.BytesIO(img_data)
            except Exception:
                return None

        # ==========================================
        # GENERACIÓN DEL ARCHIVO PDF (Compulsorio para la Biblioteca)
        # ==========================================
        doc_pdf = SimpleDocTemplate(str(ruta_pdf), pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
        styles = getSampleStyleSheet()
        
        # Estilos Customizados
        title_style = ParagraphStyle(name='PortTitle', fontName='Helvetica-Bold', fontSize=28, leading=34, alignment=1, textColor=colors.HexColor('#1a237e'))
        heading_style = ParagraphStyle(name='SecHeading', fontName='Helvetica-Bold', fontSize=18, leading=22, textColor=colors.HexColor('#1565c0'), spaceAfter=10)
        sub_heading = ParagraphStyle(name='SubHeading', fontName='Helvetica-Bold', fontSize=14, leading=18, textColor=colors.HexColor('#37474f'), spaceAfter=8)
        body_style = ParagraphStyle(name='ReportBody', fontName='Helvetica', fontSize=10, leading=14, textColor=colors.HexColor('#212121'))
        
        story = []

        # --- PÁGINA 1: PORTADA ---
        # Logo de la empresa (Ubicación relativa o placeholder) y Logo de Cliente lado a lado
        img_empresa_logo = None # Si tienes logo SSEI local puedes apuntar a él
        img_cliente_portada = None
        
        img_cliente_bytes = procesar_b64_img(cliente_logo_b64)
        if img_cliente_bytes:
            img_cliente_portada = RLImage(img_cliente_bytes, width=2.5*inch, height=1.5*inch)
            
        story.append(Spacer(1, 20))
        # Crear tabla para encabezado lado a lado de marca corporativa
        headers_data = [["LOGO CORPORATIVO SSEI", img_cliente_portada if img_cliente_portada else ""]]
        header_table = Table(headers_data, colWidths=[3.5*inch, 3.5*inch])
        header_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('ALIGN', (0,0), (0,0), 'LEFT'),
            ('ALIGN', (1,0), (1,0), 'RIGHT'),
        ]))
        story.append(header_table)
        
        story.append(Spacer(1, 150))
        story.append(Paragraph("INFORME TÉCNICO", title_style))
        story.append(Paragraph(f"ATM / PUNTO: {dto_atm}", ParagraphStyle(name='PortSubtitle', fontName='Helvetica-Bold', fontSize=16, alignment=1, textColor=colors.HexColor('#455a64'))))
        
        story.append(Spacer(1, 150))
        story.append(Paragraph(f"<b>Entidad Bancaria:</b> {dto_banco}", body_style))
        story.append(Paragraph(f"<b>Fecha de Emisión:</b> {ahora.strftime('%d-%m-%Y %H:%M')}", body_style))
        story.append(Paragraph(f"<b>Generado por:</b> SSEI Coordinación", body_style))
        story.append(PageBreak())

        # --- PÁGINA 2: SERVICIO Y SOLICITUD ---
        story.append(Paragraph("Detalle de la Solicitud de Servicio", heading_style))
        story.append(Spacer(1, 10))
        story.append(Paragraph(f"<b>Tipo de Solicitud Ingresada:</b>", sub_heading))
        story.append(Paragraph(solicitud, body_style))
        story.append(Spacer(1, 15))
        
        # Imagen de Servicio
        img_serv_bytes = procesar_b64_img(servicio_img_b64)
        if img_serv_bytes:
            story.append(Paragraph("<b>Imagen de Servicio de Origen / Evidencia de Programación</b>", sub_heading))
            story.append(Spacer(1, 5))
            r_img_serv = RLImage(img_serv_bytes, width=5.5*inch, height=3.5*inch)
            story.append(r_img_serv)
            story.append(Spacer(1, 20))
            
        story.append(PageBreak())

        # --- PÁGINA 3: DETALLE TÉCNICO PASO A PASO ---
        story.append(Paragraph("Detalle Técnico de Ejecución", heading_style))
        story.append(Paragraph("A continuación se describe la bitácora secuencial de los trabajos de ingeniería realizados en terreno:", body_style))
        story.append(Spacer(1, 15))
        
        if detalle_tecnico_pasos:
            for i, paso in enumerate(detalle_tecnico_pasos):
                paso_con_formato = f"<b>Paso {i+1}:</b> {paso}"
                story.append(Paragraph(paso_con_formato, body_style))
                story.append(Spacer(1, 8))
        else:
            story.append(Paragraph("No se registraron pasos o bitácora técnica para este servicio.", body_style))
            
        story.append(Spacer(1, 20))
        
        if comentario_general:
            story.append(Paragraph("<b>Observación General del Coordinador</b>", sub_heading))
            story.append(Paragraph(comentario_general, body_style))

        story.append(PageBreak())

        # --- PÁGINA 4: INFORME FOTOGRÁFICO ---
        story.append(Paragraph("Informe Fotográfico Estructurado", heading_style))
        story.append(Spacer(1, 10))
        
        if secciones_fotos:
            for idx, seccion in enumerate(secciones_fotos):
                titulo_foto = seccion.get("titulo", f"Fotografía de Campo {idx + 1}")
                img_b64 = seccion.get("imagen_base64")
                
                img_bytes = procesar_b64_img(img_b64)
                if img_bytes:
                    elements_to_keep = []
                    elements_to_keep.append(Paragraph(f"<b>Figura {idx+1}: {titulo_foto}</b>", sub_heading))
                    elements_to_keep.append(Spacer(1, 5))
                    r_img = RLImage(img_bytes, width=4.5*inch, height=3*inch)
                    elements_to_keep.append(r_img)
                    elements_to_keep.append(Spacer(1, 15))
                    story.append(KeepTogether(elements_to_keep))
        else:
            story.append(Paragraph("No se asociaron fotografías coordinadas de terreno al informe.", body_style))
            
        story.append(PageBreak())

        # --- PÁGINA 5: RESPALDO ORDEN DE TRABAJO (OT) ---
        story.append(Paragraph("Anexo: Validación y Cierre de Orden de Trabajo", heading_style))
        story.append(Paragraph("A continuación se adjunta copia de respaldo físico u hoja de ruta firmada en terreno conformando la solicitud:", body_style))
        story.append(Spacer(1, 15))
        
        img_ot_bytes = procesar_b64_img(ot_img_b64)
        if img_ot_bytes:
            r_img_ot = RLImage(img_ot_bytes, width=5*inch, height=5.5*inch)
            story.append(r_img_ot)
        else:
            story.append(Paragraph("<i>No se adjuntó imagen del respaldo firmado de la Orden de Trabajo.</i>", body_style))
            
        # Ejecutar compilación del PDF
        doc_pdf.build(story)
        
        # Registrar copia en la Base de Datos de la Biblioteca as `informes`
        # para que se conserve en el repositorio de documentos interactivos
        peso_bytes = ruta_pdf.stat().st_size
        ruta_relativa_sistema = f"{dto_atm}/{anio}/{mes}/{nombre_pdf}"
        
        documento_repository.create(
            nombre_original=nombre_pdf,
            nombre_sistema=ruta_relativa_sistema,
            peso_bytes=peso_bytes,
            mimetype="application/pdf",
            categoria="informes",
            banco=dto_banco,
            numero_atm=dto_atm,
            subido_por_id=token_data.get("id")
        )
        
        # ==========================================
        # CONTROL DE RETORNO SEGÚN FORMATO SELECCIONADO
        # ==========================================
        if formato == "docx":
            # Fabricamos dinámicamente un archivo editable DOCX estructurado
            doc_docx = DocxDocument()
            doc_docx.add_heading("INFORME TÉCNICO", 0)
            doc_docx.add_paragraph(f"ATM / PUNTO: {dto_atm}")
            doc_docx.add_paragraph(f"Cliente: {dto_banco}")
            doc_docx.add_paragraph(f"Fecha de Creación: {ahora.strftime('%d-%m-%Y %H:%M')}")
            doc_docx.add_paragraph(f"Tipo de Solicitud: {solicitud}")
            
            # Agregar Imagen Servicio al DOCX
            if img_serv_bytes:
                img_serv_bytes.seek(0)
                doc_docx.add_paragraph("IMAGEN DE SERVICIO")
                doc_docx.add_picture(img_serv_bytes, width=DocxInches(4.5))
                
            # Agregar Detalle Técnico Paso a Paso
            doc_docx.add_heading("Detalle Técnico de Ejecución", level=1)
            for idx, paso in enumerate(detalle_tecnico_pasos):
                doc_docx.add_paragraph(f"{idx+1}. {paso}")
                
            # Agregar Fotos del Trabajo
            doc_docx.add_heading("Registro Fotográfico de Soporte", level=1)
            for seccion in secciones_fotos:
                titulo_foto = seccion.get("titulo")
                img_b64 = seccion.get("imagen_base64")
                img_sec_bytes = procesar_b64_img(img_b64)
                if img_sec_bytes:
                    doc_docx.add_heading(titulo_foto, level=2)
                    img_sec_bytes.seek(0)
                    doc_docx.add_picture(img_sec_bytes, width=DocxInches(4.5))
                    
            # Guardamos el word
            doc_docx.save(str(ruta_docx))
            return FileResponse(path=str(ruta_docx), filename=nombre_docx, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
        else:
            return FileResponse(path=str(ruta_pdf), filename=nombre_pdf, media_type="application/pdf")
            
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Fallo en procesamiento de Reporte: {str(exc)}")